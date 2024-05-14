#!/usr/bin/env python
# coding: utf-8

# ### What does it do:
# ##### 1. Read tables from DOCX files
# ##### 2. Extract highlighted text and shading colors
# ##### 3. write the extracted data into a CSV file

# In[ ]:


import os
from docx import Document
import csv

def extract_highlight_color(run):
    # Extract highlighting color from run's XML
    xml_string = run._element.xml
    start_index = xml_string.find('<w:highlight')

    if start_index != -1:
        end_index = xml_string.find('/>', start_index) + 2
        highlight_xml = xml_string[start_index:end_index]

        color_start = highlight_xml.find('w:val="') + 7
        color_end = highlight_xml.find('"', color_start)
        color_value = highlight_xml[color_start:color_end]

        return color_value

    return None

def extract_and_format_table(docx_file):
    # Load the DOCX document
    doc = Document(docx_file)

    # Create a list to store the formatted data
    formatted_data = []

    # Iterate through tables in the document
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # Extract text from cell
                cell_text = cell.text.strip()

                # Check for highlighted text (text color)
                text_color = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.color and run.font.color.rgb:
                            text_color = run.font.color.rgb

                # Check for shading color in each cell
                shading_color = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        highlight_color = extract_highlight_color(run)
                        if highlight_color:
                            shading_color = highlight_color
                            break
                    if shading_color:
                        break

                # Append the formatted data to the list
                formatted_data.append({
                    "column_index": j,
                    "row_index": i,
                    "column_value": cell_text,
                    "text_color": text_color,
                    "shading_color": shading_color
                })

    # Extract the file name without extension
    file_name = os.path.splitext(os.path.basename(docx_file))[0]

    # Write the formatted data to a CSV file with the same name as the DOCX file
    csv_file_name = f'{file_name}.csv'
    with open(csv_file_name, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ["column_index", "row_index", "column_value", "text_color", "shading_color"]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()  # Write header to the CSV file
        writer.writerows(formatted_data)

folder_path = 'files'

# Iterate over DOCX files in the folder
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        extract_and_format_table(file_path)


# In[1]:


import os
import nbformat
from nbconvert import PythonExporter

def convert_ipynb_to_py(ipynb_file):
    # Read the IPython Notebook
    with open(ipynb_file, 'r', encoding='utf-8') as f:
        nb = nbformat.read(f, as_version=4)

    # Convert the notebook to a Python script
    py_exporter = PythonExporter()
    (py_code, _) = py_exporter.from_notebook_node(nb)

    # Write the Python script to a file
    py_file = os.path.splitext(ipynb_file)[0] + '.py'
    with open(py_file, 'w', encoding='utf-8') as f:
        f.write(py_code)

    print(f"IPython Notebook '{ipynb_file}' converted to Python script '{py_file}'.")

if __name__ == "__main__":
    ipynb_file = input("Enter the path to the IPython Notebook (.ipynb) file: ")
    convert_ipynb_to_py(ipynb_file)


# In[ ]:




